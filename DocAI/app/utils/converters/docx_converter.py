"""
DOCX file converter implementation.
Handles Microsoft Word documents.
"""
from pathlib import Path
from typing import Optional, List, Dict, Any
import base64
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.converters.base import DocumentConverter, ConversionResult
from app.models.document import DocumentMetadata
from app.core.exceptions import DocumentProcessingError


class DocxConverter(DocumentConverter):
    """
    Converter for DOCX/DOC files.
    """
    
    async def convert(self, file_path: Path) -> ConversionResult:
        """
        Convert DOCX file to standard format.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ConversionResult
        """
        if not await self.validate_file(file_path):
            return ConversionResult(
                error=f"Invalid file: {file_path}"
            )
        
        try:
            # Check if we should use LibreOffice for better formatting
            if self.config.storage.upload_folder.parent.joinpath(
                'libreoffice_uno_converter_improved.py'
            ).exists():
                # Try LibreOffice conversion first
                result = await self._convert_with_libreoffice(file_path)
                if result and not result.error:
                    return result
            
            # Fallback to python-docx
            return await self._convert_with_python_docx(file_path)
            
        except Exception as e:
            self.logger.error(
                f"Error converting DOCX file {file_path}: {e}",
                exc_info=True
            )
            return ConversionResult(
                error=f"Conversion failed: {str(e)}"
            )
    
    async def _convert_with_python_docx(self, file_path: Path) -> ConversionResult:
        """Convert using python-docx library."""
        try:
            doc = Document(str(file_path))
            
            # Extract text content
            full_text = []
            pages = []
            current_page = []
            char_count = 0
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                    current_page.append(para.text)
                    char_count += len(para.text)
                    
                    # Create new page every ~3000 characters
                    if char_count > 3000:
                        pages.append('\n'.join(current_page))
                        current_page = []
                        char_count = 0
            
            # Add remaining content
            if current_page:
                pages.append('\n'.join(current_page))
            
            # Extract metadata
            metadata = await self.extract_metadata(file_path)
            metadata.page_count = len(pages)
            metadata.word_count = len(' '.join(full_text).split())
            
            # Extract images
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        image_format = image_part.content_type.split('/')[-1]
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        images.append({
                            'format': image_format,
                            'data': image_base64,
                            'size': len(image_data)
                        })
                    except Exception as e:
                        self.logger.warning(f"Failed to extract image: {e}")
            
            return ConversionResult(
                content='\n'.join(full_text),
                pages=pages,
                content_type="text/plain",
                metadata=metadata,
                formatting_preserved=False,
                images=images if images else None
            )
            
        except Exception as e:
            self.logger.error(f"Python-docx conversion failed: {e}")
            return ConversionResult(
                error=f"Python-docx conversion failed: {str(e)}"
            )
    
    async def _convert_with_libreoffice(self, file_path: Path) -> Optional[ConversionResult]:
        """Convert using LibreOffice UNO API."""
        try:
            # Import the LibreOffice converter
            import sys
            sys.path.insert(0, str(self.config.storage.upload_folder.parent))
            from libreoffice_uno_converter_improved import render_document_with_uno_images
            
            # Convert document
            uno_result = render_document_with_uno_images(str(file_path))
            
            if uno_result and uno_result.get('success'):
                content = uno_result.get('content', '')
                
                # Extract metadata
                metadata = await self.extract_metadata(file_path)
                
                return ConversionResult(
                    content=content,
                    pages=[content],  # Single page for now
                    content_type="text/html",
                    metadata=metadata,
                    formatting_preserved=True,
                    images=uno_result.get('images')
                )
            
            return None
            
        except ImportError:
            self.logger.debug("LibreOffice converter not available")
            return None
        except Exception as e:
            self.logger.error(f"LibreOffice conversion failed: {e}")
            return None
    
    async def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """
        Extract metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            DocumentMetadata
        """
        metadata = self._extract_basic_metadata(file_path)
        
        try:
            doc = Document(str(file_path))
            
            # Extract core properties
            core_props = doc.core_properties
            
            if core_props.title:
                metadata.title = core_props.title
            if core_props.author:
                metadata.author = core_props.author
            if core_props.created:
                metadata.created_date = core_props.created
            if core_props.modified:
                metadata.modified_date = core_props.modified
            if core_props.keywords:
                metadata.keywords = [k.strip() for k in core_props.keywords.split(',')]
            
            # Count pages (estimate based on paragraphs)
            para_count = len(doc.paragraphs)
            metadata.page_count = max(1, para_count // 25)  # ~25 paragraphs per page
            
        except Exception as e:
            self.logger.warning(f"Error extracting DOCX metadata: {e}")
        
        return metadata
    
    async def edit_document(
        self,
        file_path: Path,
        operation: str,
        target: str,
        replacement: Optional[str] = None,
        preserve_formatting: bool = True
    ) -> bool:
        """
        Edit DOCX document content.
        
        Args:
            file_path: Path to document
            operation: Edit operation ('replace', 'insert', 'delete')
            target: Target text
            replacement: Replacement text
            preserve_formatting: Whether to preserve formatting
            
        Returns:
            True if successful
        """
        try:
            doc = Document(str(file_path))
            changes_made = False
            
            if operation == 'replace':
                # Find and replace text
                for para in doc.paragraphs:
                    if target in para.text:
                        if preserve_formatting:
                            # Preserve run formatting
                            for run in para.runs:
                                if target in run.text:
                                    run.text = run.text.replace(target, replacement or '')
                                    changes_made = True
                        else:
                            # Simple text replacement
                            para.text = para.text.replace(target, replacement or '')
                            changes_made = True
                
                # Also check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if target in para.text:
                                    para.text = para.text.replace(target, replacement or '')
                                    changes_made = True
            
            if changes_made:
                # Save the document
                doc.save(str(file_path))
                self.logger.info(f"Document edited successfully: {file_path}")
                return True
            else:
                self.logger.warning(f"No changes made - target text not found: {target}")
                return False
                
        except Exception as e:
            self.logger.error(f"Error editing DOCX: {e}", exc_info=True)
            return False
    
    async def export(
        self,
        source_path: Path,
        target_path: Path,
        format: str,
        include_metadata: bool = False
    ) -> bool:
        """
        Export DOCX to different format.
        
        Args:
            source_path: Source document path
            target_path: Target path
            format: Target format
            include_metadata: Whether to include metadata
            
        Returns:
            True if successful
        """
        try:
            if format == 'docx':
                # Just copy the file
                import shutil
                shutil.copy2(source_path, target_path)
                return True
            
            elif format == 'txt':
                # Convert to plain text
                result = await self.convert(source_path)
                if result.content:
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(result.content)
                    return True
            
            elif format == 'html':
                # Try LibreOffice conversion
                result = await self._convert_with_libreoffice(source_path)
                if result and result.content:
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(result.content)
                    return True
            
            self.logger.warning(f"Export format not supported: {format}")
            return False
            
        except Exception as e:
            self.logger.error(f"Export failed: {e}", exc_info=True)
            return False