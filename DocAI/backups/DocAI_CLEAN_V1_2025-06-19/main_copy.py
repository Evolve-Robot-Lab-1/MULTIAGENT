import os
import sys
import logging
import tempfile
import time
import base64
from io import BytesIO

# Set environment variables before any other imports
os.environ['POCKETGROQ_DISABLE_AUTO_RAG'] = 'true'
os.environ['LANGCHAIN_USER_AGENT'] = 'durga-ai-doc-agent/1.0'
os.environ['TORCH_LOAD_WEIGHTS_ONLY'] = 'True'

# Enhanced logging configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8', mode='a'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Add file handler for errors specifically
error_handler = logging.FileHandler('error.log', encoding='utf-8', mode='a')
error_handler.setLevel(logging.ERROR)
error_handler.setFormatter(logging.Formatter(
    '%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'
))
logger.addHandler(error_handler)

logger.info("Environment variables set")

# Now import everything else
from flask import Flask, request, Response, send_from_directory, jsonify, send_file
from flask_cors import CORS
import groq
from dotenv import load_dotenv
import json
import urllib.parse
from werkzeug.utils import secure_filename
# import google.generativeai as genai
from flask import session
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re

# Import local modules
from pocketgroq.groq_provider import GroqProvider
from rag_handler import RAGHandler
# from libreoffice_converter import render_document_with_libreoffice
# from libreoffice_uno_converter import render_document_with_uno

# Load environment variables
load_dotenv()
logger.info("Environment variables loaded")

# Initialize Flask app
app = Flask(__name__, static_folder='static2.0')
CORS(app, supports_credentials=True)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "fallback_secret_key")

# Add configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
DOCUMENTS_FOLDER = os.path.join(UPLOAD_FOLDER, 'documents')

# Configure app
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Global error handlers
@app.errorhandler(404)
def not_found_error(error):
    logger.error(f"404 Error: {request.url}")
    return jsonify({
        'error': 'Resource not found',
        'url': request.url
    }), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500 Error: {str(error)}", exc_info=True)
    return jsonify({
        'error': 'Internal server error',
        'message': str(error) if app.debug else 'An unexpected error occurred'
    }), 500

@app.errorhandler(Exception)
def handle_exception(error):
    logger.error(f"Unhandled Exception: {str(error)}", exc_info=True)
    return jsonify({
        'error': 'Server error',
        'message': str(error) if app.debug else 'An unexpected error occurred'
    }), 500

# Initialize core services with shared provider
groq_api_key = os.environ.get("GROQ_API_KEY")
groq_client = groq.Groq(api_key=groq_api_key)

# Create a single GroqProvider instance
groq_provider = GroqProvider(api_key=groq_api_key)
os.environ["GROQ_DEFAULT_MODEL"] = "llama3-groq-70b-8192-tool-use-preview"  # Set default model
logger.info("Core services initialized with llama3-groq-70b-8192-tool-use-preview model")

# Initialize optional services
# genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
# OpenAI client removed - using only GROQ

# Initialize caches and handlers
logger.info("Core handlers initialized")

# Initialize RAG handler
rag_handler = RAGHandler(uploads_dir=app.config['UPLOAD_FOLDER'])
logger.info("RAG handler initialized")

# Whisper model removed - audio transcription functionality removed

# Initialize chat histories
app.chat_histories = {}

def get_chat_history():
    if 'chat_id' not in session:
        session['chat_id'] = str(uuid.uuid4())
    
    chat_id = session['chat_id']
    if chat_id not in app.chat_histories:
        app.chat_histories[chat_id] = []
    
    # Always use English - system message
    system_message = {
        "role": "system",
        "content": "You are Durga AI, a helpful document processing assistant. Please respond in English."
    }
    
    if not app.chat_histories[chat_id] or app.chat_histories[chat_id][0]["role"] != "system":
        app.chat_histories[chat_id].insert(0, system_message)
    elif app.chat_histories[chat_id][0]["role"] == "system":
        app.chat_histories[chat_id][0] = system_message
    
    return app.chat_histories[chat_id]

def create_prompt(user_input):
    return f"You are Durga AI, a helpful document processing assistant. Respond to the following in English: {user_input}"

@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'indexf.html')

@app.route('/static2.0/<path:filename>')
def serve_static(filename):
    """Serve static files from static2.0 folder"""
    return send_from_directory(app.static_folder, filename)

# Translation function removed - focusing on core document functionality

@app.route('/api/simple_chat', methods=['POST'])
def handle_simple_chat():
    """Simple chat endpoint for non-streaming responses"""
    try:
        data = request.json
        message = data.get('message', '')
        language = data.get('language', 'en')
        
        if not message:
            return jsonify({'status': 'error', 'response': 'No message provided'}), 400
        
        chat_history = get_chat_history()
        chat_history.append({"role": "user", "content": message})
        
        # Use the same model as streaming endpoint
        model_to_use = "llama-3.3-70b-versatile"
        
        # Get non-streaming response
        response = groq_client.chat.completions.create(
            messages=chat_history,
            model=model_to_use,
            max_tokens=2000,
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content
        chat_history.append({"role": "assistant", "content": ai_response})
        
        return jsonify({
            'status': 'success',
            'response': ai_response
        })
        
    except Exception as e:
        logger.error(f"Error in simple_chat: {str(e)}")
        return jsonify({
            'status': 'error',
            'response': 'An error occurred while processing your request.'
        }), 500

@app.route('/api/query_stream', methods=['POST'])
def handle_query_stream():
    data = request.json
    logging.info(f"Received request: {data}")
    user_input = data.get('query')
    selected_model = data.get('model')
    use_rag = data.get('use_rag', True)  # Default to using RAG if available
    language = 'en'  # Fixed to English
    
    chat_history = get_chat_history()
    chat_history.append({"role": "user", "content": user_input})
    
    logging.info(f"Selected model: {selected_model}")
    # Choose model based on selection
    if selected_model == "deepseek":
        model_to_use = "deepseek-r1-distill-llama-70b"
    else:
        model_to_use = "llama-3.3-70b-versatile"

    def generate():
        try:
            # Check if we have documents and should use RAG
            if use_rag and rag_handler.has_documents():
                from langchain_groq import ChatGroq
                llm = ChatGroq(
                    api_key=os.environ.get("GROQ_API_KEY"),
                    model_name=model_to_use
                )
                
                # Use user input directly for RAG (translation removed for simplicity)
                search_query = user_input
                
                # Get RAG response
                rag_response = rag_handler.query_documents(search_query, llm)
                if rag_response:
                    # Stream the RAG response
                    for char in rag_response:
                        yield f"data: {json.dumps({'content': char})}\n\n"
                    yield "data: [DONE]\n\n"
                    chat_history.append({"role": "assistant", "content": rag_response})
                    return

            # If no RAG response or no documents, proceed with regular chat
            stream = groq_client.chat.completions.create(
                messages=chat_history,
                model=model_to_use,
                max_tokens=5000,
                temperature=0.7,
                top_p=0.95,
                stream=True
            )

            full_response = ""
            for chunk in stream:
                if chunk and chunk.choices and chunk.choices[0].delta:
                    if chunk.choices[0].delta.content is not None:
                        content = chunk.choices[0].delta.content
                        unescaped = content.encode('utf-8').decode('unicode-escape')
                        url_decoded = urllib.parse.unquote(unescaped)
                        decoded_content = url_decoded.encode('raw_unicode_escape').decode('utf-8')
                        full_response += decoded_content
                        yield f"data: {json.dumps(decoded_content)}\n\n"
            
            chat_history.append({"role": "assistant", "content": full_response})
            logging.info(f"Chat history updated. Current length: {len(chat_history)}")

        except Exception as e:
            logging.error(f"Error in generate(): {str(e)}")
            error_message = "An error occurred while processing your request."
            yield f"data: {json.dumps(error_message)}\n\n"

    return Response(generate(), mimetype='text/event-stream')

# Audio transcription functionality removed - focusing on document functionality

# Duplicate route removed - using static2.0 route above

# Web search tool routes removed - focusing on document functionality

# Add this new route for clearing chat history
@app.route('/api/clear_chat', methods=['POST'])
def clear_chat():
    if 'chat_id' in session:
        chat_id = session['chat_id']
        if chat_id in app.chat_histories:
            app.chat_histories[chat_id] = [
                {"role": "system", "content": "You are Durga AI, a helpful document processing assistant. Please respond in English."}
            ]
    return jsonify({"message": "Chat history cleared"}), 200

# Update the clear_memory function
@app.route('/api/clear_memory', methods=['POST'])
def clear_memory():
    logging.info("Clear memory request received")
    if 'chat_id' in session:
        chat_id = session['chat_id']
        if chat_id in app.chat_histories:
            del app.chat_histories[chat_id]
        session.pop('chat_id')
        logging.info(f"Cleared chat_id {chat_id} from session and removed chat history")
    else:
        logging.info("No chat_id found in session to clear")
    return jsonify({"message": "User memory cleared"}), 200

# Scraping functionality removed - focusing on document functionality

# Web search API route removed - focusing on document functionality

# Language selection route removed - using English only

# Agent functionality removed - focusing on document functionality

# Writer functionality removed - focusing on document functionality

# Initialize RAG handler
rag_handler = RAGHandler()

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/rag/upload', methods=['POST'])
def upload_document():
    try:
        if 'document' not in request.files:
            logger.warning("Upload attempt with no document part")
            return jsonify({'error': 'No document part'}), 400
        
        file = request.files['document']
        if file.filename == '':
            logger.warning("Upload attempt with empty filename")
            return jsonify({'error': 'No selected file'}), 400
        
        if not allowed_file(file.filename):
            logger.warning(f"Upload attempt with invalid file type: {file.filename}")
            return jsonify({'error': f'Invalid file type. Allowed types are: {", ".join(ALLOWED_EXTENSIONS)}'}), 400
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        logger.info(f"Processing upload for file: {filename}")
        
        try:
            file.save(file_path)
            logger.info(f"File saved successfully: {file_path}")
            
            success = rag_handler.process_document(file_path, filename)
            if success:
                doc_count = len(rag_handler.document_store)
                logger.info(f"Document processed successfully. Total documents: {doc_count}")
                return jsonify({
                    'message': 'Document processed successfully',
                    'document_count': doc_count,
                    'filename': filename  # Return the secure filename
                })
            else:
                logger.error(f"Failed to process document: {filename}")
                return jsonify({'error': 'Failed to process document'}), 500
                
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}", exc_info=True)
            return jsonify({'error': f'Error processing document: {str(e)}'}), 500
        finally:
            # Clean up temporary file
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info(f"Temporary file removed: {file_path}")
                except Exception as e:
                    logger.error(f"Error removing temporary file {file_path}: {str(e)}")
    
    except Exception as e:
        logger.error(f"Unexpected error in upload_document: {str(e)}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred'}), 500

@app.route('/rag/remove', methods=['POST'])
def remove_document():
    data = request.get_json()
    if not data or 'filename' not in data:
        return jsonify({'error': 'No filename provided'}), 400
    
    filename = data['filename']
    success = rag_handler.remove_document(filename)
    
    if success:
        return jsonify({
            'message': 'Document removed successfully',
            'document_count': len(rag_handler.document_store)
        })
    else:
        return jsonify({'error': 'Failed to remove document'}), 500

@app.route('/rag/status', methods=['GET'])
def get_rag_status():
    return jsonify(rag_handler.get_status())

@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    if not data or 'message' not in data:
        return jsonify({'error': 'No message provided'}), 400
    
    message = data['message']
    use_rag = data.get('use_rag', False)
    
    try:
        if use_rag and rag_handler.has_documents():
            # Create LLM instance for RAG
            from langchain_groq import ChatGroq
            llm = ChatGroq(
                api_key=os.environ.get("GROQ_API_KEY"),
                model_name="llama-3.3-70b-versatile"
            )
            response = rag_handler.query_documents(message, llm)
            if response is None:
                response = "I couldn't find relevant information in the documents. Please try rephrasing your question."
        else:
            # Use regular chat completion
            response = "Regular chat completion response"  # Replace with your chat completion logic
        
        return jsonify({'response': response})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    uploaded_files = []
    
    try:
        # Ensure upload directories exist
        os.makedirs(DOCUMENTS_FOLDER, exist_ok=True)
        
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                # Save to documents directory
                document_path = os.path.join(DOCUMENTS_FOLDER, filename)
                
                # Save original file
                file.save(document_path)
                uploaded_files.append(filename)
                
                logger.info(f"File saved to: {document_path}")
                
                # Process based on file type
                file_ext = filename.rsplit('.', 1)[1].lower()
                if file_ext == 'pdf':
                    # Process PDF for RAG while keeping original
                    rag_handler.process_document(document_path, filename)
                elif file_ext in ['doc', 'docx']:
                    # Process for RAG
                    rag_handler.process_document(document_path, filename)

        return jsonify({
            'message': 'Files uploaded successfully',
            'files': uploaded_files,
            'total_documents': len(uploaded_files)
        })

    except Exception as e:
        logger.error(f"Error uploading files: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Debug function removed - not needed

@app.route('/view_document/<path:filename>')
def view_document(filename):
    """View a document with preserved formatting"""
    try:
        document_path = os.path.join(DOCUMENTS_FOLDER, filename)
        if not os.path.exists(document_path):
            return jsonify({'error': 'Document not found'}), 404

        file_ext = os.path.splitext(filename)[1].lower()
        
        # For PDF files, serve directly
        if file_ext == '.pdf':
            return send_file(
                document_path,
                mimetype='application/pdf',
                as_attachment=False,
                download_name=filename
            )
        
        # For Word and text documents, process with formatting preservation
        return process_text_document(document_path)
    
    except Exception as e:
        logger.error(f"Error viewing docment: {str(e)}")
        return jsonify({'error': str(e)}), 500

def handle_text_document(document_path: str, file_ext: str):
    """Handle text-based documents"""
    try:
        pages = []
        if file_ext in ['doc', 'docx']:
            import docx
            doc = docx.Document(document_path)
            current_page = []
            char_count = 0
            
            for para in doc.paragraphs:
                if para.text.strip():
                    current_page.append(para.text)
                    char_count += len(para.text)
                    if char_count > 3000:
                        pages.append('<br>'.join(current_page))
                        current_page = []
                        char_count = 0
            
            if current_page:
                pages.append('<br>'.join(current_page))
                
        elif file_ext == 'txt':
            with open(document_path, 'r', encoding='utf-8') as f:
                content = f.read()
                pages = [content[i:i+3000].replace('\n', '<br>') 
                        for i in range(0, len(content), 3000)]

        return jsonify({
            'pages': pages or ['No content found in document'],
            'total_pages': len(pages)
        })

    except Exception as e:
        logger.error(f"Error processing text document: {str(e)}")
        raise

@app.route('/convert-to-word', methods=['POST'])
def convert_to_word():
    try:
        data = request.get_json()
        filename = data.get('filename')
        
        if not filename:
            return jsonify({'error': 'No filename provided'}), 400
        
        document_path = os.path.join(DOCUMENTS_FOLDER, filename)
        logger.info(f"Looking for document at: {document_path}")
        
        if not os.path.exists(document_path):
            return jsonify({
                'error': 'Document not found',
                'path': document_path,
                'exists': os.path.exists(document_path)
            }), 404
        
        # Simple conversion logic without external dependencies
        if filename.lower().endswith('.pdf'):
            # Return the same file for now - conversion functionality removed
            return jsonify({
                'success': True,
                'message': 'PDF ready for viewing',
                'converted_file': filename
            })
            
            return jsonify({
                'success': True,
            'message': 'Document ready for viewing',
            'converted_file': filename
            })
            
    except Exception as e:
        logger.error(f"Error in convert_to_word: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/edit-document', methods=['POST'])
def edit_document():
    """Handle document edits"""
    try:
        data = request.get_json()
        filename = data.get('filename')
        original_text = data.get('originalText')
        new_text = data.get('newText')
        
        if not all([filename, original_text, new_text]):
            return jsonify({'error': 'Missing required parameters'}), 400
        
        document_path = os.path.join(DOCUMENTS_FOLDER, filename)
        if not os.path.exists(document_path):
            return jsonify({'error': 'Document not found'}), 404
        
        # Handle Word document edits
        if filename.endswith(('.doc', '.docx')):
            doc = Document(document_path)
            
            # Find and replace text while preserving formatting
            for paragraph in doc.paragraphs:
                if original_text in paragraph.text:
                    # Store original formatting
                    original_runs = []
                    for run in paragraph.runs:
                        if original_text in run.text:
                            original_runs.append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font': run.font.name,
                                'size': run.font.size
                            })
                    
                    # Clear paragraph
                    paragraph.clear()
                    
                    # Add new text with original formatting
                    for original_run in original_runs:
                        new_run = paragraph.add_run(original_run['text'].replace(original_text, new_text))
                        new_run.bold = original_run['bold']
                        new_run.italic = original_run['italic']
                        new_run.underline = original_run['underline']
                        new_run.font.name = original_run['font']
                        if original_run['size']:
                            new_run.font.size = original_run['size']
            
            # Save the changes
            doc.save(document_path)
            
            return jsonify({
                'success': True,
                'message': 'Edit applied successfully'
            })
        
        return jsonify({'error': 'Unsupported file type for editing'}), 400
        
    except Exception as e:
        logger.error(f"Error in edit_document: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

def process_text_document(document_path):
    try:
        file_ext = os.path.splitext(document_path)[1].lower()
        pages = []
        
        if file_ext in ['.doc', '.docx']:
            # Use improved LibreOffice UNO API converter ONLY (no fallbacks)
            logger.info("🚀 Attempting improved LibreOffice UNO API conversion...")
            try:
                from libreoffice_uno_converter_improved import render_document_with_uno_images
                uno_result = render_document_with_uno_images(document_path)
                
                if uno_result and uno_result.get('success'):
                    logger.info(f"✅ USING IMPROVED LIBREOFFICE UNO API - {uno_result.get('images_found', 0)} images embedded")
                    # Convert content to pages format expected by frontend
                    content = uno_result.get('content', '')
                    pages = [content] if content else []
                    return jsonify({
                        'pages': pages,
                        'total_pages': len(pages),
                        'success': True,
                        'method': uno_result.get('method'),
                        'images_found': uno_result.get('images_found', 0)
                    })
                else:
                    error_msg = uno_result.get('error', 'Unknown error') if uno_result else 'Converter returned None'
                    logger.error(f"❌ LibreOffice UNO API conversion failed: {error_msg}")
                    return jsonify({
                        'error': f'Document conversion failed: {error_msg}',
                        'success': False,
                        'details': uno_result
                    }), 500
                    
            except ImportError as e:
                logger.error(f"❌ Failed to import improved UNO converter: {str(e)}")
                return jsonify({
                    'error': 'LibreOffice UNO converter not available. Please check installation.',
                    'success': False
                }), 500
            except Exception as e:
                logger.error(f"❌ Unexpected error in UNO conversion: {str(e)}")
                return jsonify({
                    'error': f'Document conversion error: {str(e)}',
                    'success': False
                }), 500
        
        elif file_ext == '.txt':
            with open(document_path, 'r', encoding='utf-8') as f:
                content = f.read()
                pages = [f'<div class="word-content"><div class="word-page"><p style="white-space: pre-wrap; font-family: monospace; font-size: 11pt; color: #000000;">{content[i:i+3000]}</p></div></div>' 
                        for i in range(0, len(content), 3000)]
        
        return jsonify({
            'pages': pages or ['No content found in document'],
            'total_pages': len(pages)
        })

    except Exception as e:
        logger.error(f"Error processing text document: {str(e)}")
        raise

def process_paragraph(para):
    """Process a paragraph with all its formatting"""
    # Get paragraph properties
    style_parts = []
    classes = []
    
    # Font family and size from paragraph style
    if para.style:
        if hasattr(para.style, 'font'):
            if para.style.font.name:
                style_parts.append(f"font-family: '{para.style.font.name}', sans-serif")
            if para.style.font.size:
                style_parts.append(f"font-size: {para.style.font.size.pt}pt")
    
    # Paragraph alignment
    if para.alignment:
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            style_parts.append("text-align: center")
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            style_parts.append("text-align: right")
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            style_parts.append("text-align: justify")
    
    # Paragraph spacing
    if hasattr(para.paragraph_format, 'space_before') and para.paragraph_format.space_before:
        style_parts.append(f"margin-top: {para.paragraph_format.space_before.pt}pt")
    if hasattr(para.paragraph_format, 'space_after') and para.paragraph_format.space_after:
        style_parts.append(f"margin-bottom: {para.paragraph_format.space_after.pt}pt")
    
    # Line spacing
    if hasattr(para.paragraph_format, 'line_spacing') and para.paragraph_format.line_spacing:
        style_parts.append(f"line-height: {para.paragraph_format.line_spacing}")
    
    # First line indent
    if hasattr(para.paragraph_format, 'first_line_indent') and para.paragraph_format.first_line_indent:
        style_parts.append(f"text-indent: {para.paragraph_format.first_line_indent.pt}pt")
    
    # Left indent
    if hasattr(para.paragraph_format, 'left_indent') and para.paragraph_format.left_indent:
        indent_level = int(para.paragraph_format.left_indent.pt / 36)  # 36pt = 0.5in
        if indent_level > 0:
            classes.append(f"indent-{min(indent_level, 4)}")
    
    # Build paragraph HTML
    style_str = '; '.join(style_parts) if style_parts else ''
    class_str = ' '.join(classes) if classes else ''
    
    # Process runs within the paragraph
    para_html = []
    for run in para.runs:
        run_html = process_run(run)
        para_html.append(run_html)
    
    # Determine paragraph tag based on style
    tag = 'p'
    if para.style and para.style.name:
        if 'Heading 1' in para.style.name:
            tag = 'h1'
        elif 'Heading 2' in para.style.name:
            tag = 'h2'
        elif 'Heading 3' in para.style.name:
            tag = 'h3'
    
    return f'<{tag} style="{style_str}" class="{class_str}">{"".join(para_html)}</{tag}>'

def process_paragraph_with_content(para, content_list):
    """Process a paragraph with pre-processed content"""
    # Get paragraph properties
    style_parts = []
    classes = []
    
    # Get paragraph background color
    if hasattr(para, '_element') and hasattr(para.paragraph_format, 'element'):
        # Check for shading/background color in paragraph properties
        pr_element = para.paragraph_format.element
        if pr_element is not None:
            # Look for shading element
            for child in pr_element:
                if child.tag.endswith('shd'):
                    fill = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill and fill != 'auto':
                        style_parts.append(f"background-color: #{fill}")
    
    # Font family and size from paragraph style
    if para.style:
        if hasattr(para.style, 'font'):
            if para.style.font.name:
                style_parts.append(f"font-family: '{para.style.font.name}', sans-serif")
            if para.style.font.size:
                style_parts.append(f"font-size: {para.style.font.size.pt}pt")
    
    # Paragraph alignment
    if para.alignment:
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            style_parts.append("text-align: center")
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            style_parts.append("text-align: right")
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            style_parts.append("text-align: justify")
    
    # Paragraph spacing
    if hasattr(para.paragraph_format, 'space_before') and para.paragraph_format.space_before:
        style_parts.append(f"margin-top: {para.paragraph_format.space_before.pt}pt")
    if hasattr(para.paragraph_format, 'space_after') and para.paragraph_format.space_after:
        style_parts.append(f"margin-bottom: {para.paragraph_format.space_after.pt}pt")
    
    # Build paragraph HTML
    style_str = '; '.join(style_parts) if style_parts else ''
    class_str = ' '.join(classes) if classes else ''
    
    # Determine paragraph tag based on style
    tag = 'p'
    if para.style and para.style.name:
        if 'Heading 1' in para.style.name:
            tag = 'h1'
        elif 'Heading 2' in para.style.name:
            tag = 'h2'
        elif 'Heading 3' in para.style.name:
            tag = 'h3'
    
    return f'<{tag} style="{style_str}" class="{class_str}">{"".join(content_list)}</{tag}>'

def process_run(run):
    """Process a run (text with consistent formatting)"""
    if not run.text:
        return ''
    
    style_parts = []
    classes = []
    
    # Font properties
    if run.font.name:
        style_parts.append(f"font-family: '{run.font.name}', sans-serif")
    if run.font.size:
        style_parts.append(f"font-size: {run.font.size.pt}pt")
    
    # Bold
    if run.bold:
        style_parts.append("font-weight: bold")
    
    # Italic
    if run.italic:
        style_parts.append("font-style: italic")
    
    # Underline
    if run.underline:
        classes.append("text-underline")
    
    # Strike-through
    if run.font.strike:
        classes.append("text-strikethrough")
    
    # Subscript/Superscript
    if run.font.subscript:
        classes.append("text-subscript")
    elif run.font.superscript:
        classes.append("text-superscript")
    
    # Text color
    if run.font.color and run.font.color.rgb:
        rgb = run.font.color.rgb
        # Convert RGBColor to hex string
        if hasattr(rgb, '__iter__'):
            r, g, b = rgb
            style_parts.append(f"color: #{r:02x}{g:02x}{b:02x}")
        else:
            # If it's already an integer, use it directly
            style_parts.append(f"color: #{rgb:06x}")
    
    # Highlight color
    if run.font.highlight_color:
        # Map Word highlight colors to CSS
        highlight_colors = {
            1: "#000000",  # black
            2: "#0000FF",  # blue
            3: "#00FFFF",  # cyan
            4: "#00FF00",  # green
            5: "#FF00FF",  # magenta
            6: "#FF0000",  # red
            7: "#FFFF00",  # yellow
            8: "#FFFFFF",  # white
            9: "#000080",  # dark blue
            10: "#008080", # dark cyan
            11: "#008000", # dark green
            12: "#800080", # dark magenta
            13: "#800000", # dark red
            14: "#808000", # dark yellow
            15: "#808080", # dark gray
            16: "#C0C0C0"  # light gray
        }
        color_value = run.font.highlight_color
        if isinstance(color_value, int) and color_value in highlight_colors:
            style_parts.append(f"background-color: {highlight_colors[color_value]}")
        else:
            style_parts.append("background-color: yellow")
    
    # Check for text shading/background
    if hasattr(run, '_element'):
        # Look for shading in run properties
        for child in run._element:
            if child.tag.endswith('rPr'):  # Run properties
                for prop in child:
                    if prop.tag.endswith('shd'):
                        fill = prop.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if fill and fill != 'auto':
                            style_parts.append(f"background-color: #{fill}")
    
    # Build span if formatting exists
    if style_parts or classes:
        style_str = '; '.join(style_parts) if style_parts else ''
        class_str = ' '.join(classes) if classes else ''
        return f'<span style="{style_str}" class="{class_str}">{run.text}</span>'
    else:
        return run.text

def process_table(table):
    """Process a table with all its formatting"""
    html_parts = ['<table>']
    
    # Process rows
    for row in table.rows:
        html_parts.append('<tr>')
        for cell in row.cells:
            # Check if this is a header cell (usually first row)
            tag = 'th' if row == table.rows[0] else 'td'
            
            # Process cell content
            cell_content = []
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    cell_content.append(process_paragraph(paragraph))
            
            html_parts.append(f'<{tag}>{"".join(cell_content)}</{tag}>')
        html_parts.append('</tr>')
    
    html_parts.append('</table>')
    return '\n'.join(html_parts)

def extract_image_by_rel_id(rel, doc):
    """Extract image directly from relationship"""
    try:
        image_part = rel.target_part
        image_data = image_part.blob
        
        # Determine image format
        image_format = image_part.content_type.split('/')[-1]
        if image_format == 'jpeg':
            image_format = 'jpg'
        
        # Convert to base64
        image_base64 = base64.b64encode(image_data).decode('utf-8')
        
        # Create image HTML with reasonable default size
        img_html = f'''
        <div class="word-image-container" style="text-align: center; margin: 10px 0;">
            <img src="data:image/{image_format};base64,{image_base64}" 
                 style="max-width: 100%; max-height: 400px; height: auto;"
                 alt="Document image" />
        </div>
        '''
        
        return img_html
        
    except Exception as e:
        logger.error(f"Error extracting image by rel ID: {str(e)}")
        return None

def extract_image_from_run(run, doc):
    """Extract image with proper size and position from a run"""
    try:
        import xml.etree.ElementTree as ET
        
        # Parse the run's XML
        run_xml = run._element.xml
        
        # Look for drawing elements
        if 'w:drawing' in run_xml:
            # Find all namespace prefixes
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }
            
            # Extract dimensions from extent
            cx = cy = None
            inline = False
            anchor = False
            
            # Check for extent in XML
            if 'wp:extent' in run_xml:
                # Extract cx and cy values
                import re
                cx_match = re.search(r'<wp:extent[^>]*cx="(\d+)"', run_xml)
                cy_match = re.search(r'<wp:extent[^>]*cy="(\d+)"', run_xml)
                
                if cx_match:
                    cx = int(cx_match.group(1))
                if cy_match:
                    cy = int(cy_match.group(1))
            
            # Check if inline or anchored
            inline = '<wp:inline' in run_xml
            anchor = '<wp:anchor' in run_xml
            
            # Extract embed ID
            embed_match = re.search(r'r:embed="([^"]+)"', run_xml)
            if embed_match:
                embed_id = embed_match.group(1)
                
                # Get the image from relationships
                if embed_id in doc.part.rels:
                    rel = doc.part.rels[embed_id]
                    if "image" in rel.reltype:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        
                        # Determine image format
                        image_format = image_part.content_type.split('/')[-1]
                        if image_format == 'jpeg':
                            image_format = 'jpg'
                        
                        # Convert to base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        # Convert EMUs to pixels (1 inch = 914400 EMUs = 96 pixels)
                        width = int(cx / 914400 * 96) if cx else 300
                        height = int(cy / 914400 * 96) if cy else 200
                        
                        # Maintain aspect ratio
                        if cx and cy:
                            aspect_ratio = cy / cx
                            # Limit max width for better display
                            if width > 600:
                                width = 600
                                height = int(width * aspect_ratio)
                        
                        # Create appropriate HTML based on position type
                        if inline:
                            img_html = f'''<img src="data:image/{image_format};base64,{image_base64}" 
                                 style="width: {width}px; height: {height}px; vertical-align: middle; margin: 2px 5px;"
                                 alt="Document image" />'''
                        else:
                            # Check for alignment
                            align = 'center'
                            if 'positionH' in run_xml:
                                if 'align="left"' in run_xml:
                                    align = 'left'
                                elif 'align="right"' in run_xml:
                                    align = 'right'
                            
                            img_html = f'''
                            <div class="word-image-container" style="text-align: {align}; margin: 10px 0; clear: both;">
                                <img src="data:image/{image_format};base64,{image_base64}" 
                                     style="width: {width}px; height: {height}px; max-width: 100%;"
                                     alt="Document image" />
                            </div>
                            '''
                        
                        return {'html': img_html, 'inline': inline}
        
        # Also check for pict elements (legacy images)
        elif 'v:imagedata' in run_xml:
            # Extract r:id for legacy images
            import re
            id_match = re.search(r'r:id="([^"]+)"', run_xml)
            if id_match:
                rel_id = id_match.group(1)
                if rel_id in doc.part.rels:
                    rel = doc.part.rels[rel_id]
                    if "image" in rel.reltype:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        
                        # Determine image format
                        image_format = image_part.content_type.split('/')[-1]
                        if image_format == 'jpeg':
                            image_format = 'jpg'
                        
                        # Convert to base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        # Extract dimensions from shape style
                        width = height = None
                        style_match = re.search(r'style="([^"]+)"', run_xml)
                        if style_match:
                            style = style_match.group(1)
                            width_match = re.search(r'width:([0-9.]+)pt', style)
                            height_match = re.search(r'height:([0-9.]+)pt', style)
                            
                            if width_match:
                                width = int(float(width_match.group(1)) * 96 / 72)  # pt to px
                            if height_match:
                                height = int(float(height_match.group(1)) * 96 / 72)
                        
                        if not width:
                            width = 300
                        if not height:
                            height = 200
                        
                        img_html = f'''
                        <div class="word-image-container" style="text-align: center; margin: 10px 0;">
                            <img src="data:image/{image_format};base64,{image_base64}" 
                                 style="width: {width}px; height: {height}px; max-width: 100%;"
                                 alt="Document image" />
                        </div>
                        '''
                        
                        return {'html': img_html, 'inline': False}
        
        return None
        
    except Exception as e:
        logger.error(f"Error extracting image from run: {str(e)}")
        return None


@app.route('/save-document', methods=['POST'])
def save_document():
    try:
        data = request.get_json()
        filename = data.get('filename')
        content = data.get('content')
        
        if not filename:
            return jsonify({'error': 'Filename is required'}), 400
            
        # Create a new document with the content
        doc = Document()
        
        # Parse HTML content
        soup = BeautifulSoup(content, 'html.parser')
        
        # Find all elements (paragraphs, headings, tables)
        elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'table', 'ul', 'ol'])
        
        for element in elements:
            if element.name == 'table':
                # Process table
                process_html_table_to_docx(doc, element)
            elif element.name in ['ul', 'ol']:
                # Process list
                process_html_list_to_docx(doc, element)
            else:
                # Process paragraph or heading
                process_html_paragraph_to_docx(doc, element)
        
        # Save the document
        file_path = os.path.join(DOCUMENTS_FOLDER, filename)
        doc.save(file_path)
            
        return jsonify({'success': True, 'message': 'Document saved successfully', 'path': file_path})
        
    except Exception as e:
        logger.error(f"Error saving document: {str(e)}")
        return jsonify({'error': str(e)}), 500

def process_html_paragraph_to_docx(doc, element):
    """Convert HTML paragraph/heading to Word paragraph with formatting"""
    # Determine if it's a heading
    if element.name in ['h1', 'h2', 'h3']:
        para = doc.add_heading(level=int(element.name[1]))
    else:
        para = doc.add_paragraph()
    
    # Get paragraph style
    style = element.get('style', '')
    style_dict = parse_style_string(style)
    
    # Apply paragraph alignment
    if 'text-align' in style_dict:
        align = style_dict['text-align']
        if align == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Apply paragraph spacing
    if 'margin-top' in style_dict:
        para.paragraph_format.space_before = Pt(parse_pt_value(style_dict['margin-top']))
    if 'margin-bottom' in style_dict:
        para.paragraph_format.space_after = Pt(parse_pt_value(style_dict['margin-bottom']))
    if 'text-indent' in style_dict:
        para.paragraph_format.first_line_indent = Pt(parse_pt_value(style_dict['text-indent']))
    if 'line-height' in style_dict:
        try:
            para.paragraph_format.line_spacing = float(style_dict['line-height'])
        except:
            pass
    
    # Process inline content
    process_inline_content(para, element)

def process_inline_content(para, element):
    """Process inline content (text and spans) within a paragraph"""
    for child in element.children:
        if isinstance(child, str):
            # Plain text
            if child.strip():
                para.add_run(child)
        elif child.name == 'span':
            # Formatted text
            run = para.add_run(child.get_text())
            
            # Get span styles
            style = child.get('style', '')
            style_dict = parse_style_string(style)
            classes = child.get('class', [])
            if isinstance(classes, str):
                classes = classes.split()
            
            # Apply font formatting
            if 'font-family' in style_dict:
                run.font.name = style_dict['font-family'].split(',')[0].strip().replace("'", "").replace('"', '')
            if 'font-size' in style_dict:
                run.font.size = Pt(parse_pt_value(style_dict['font-size']))
            if 'font-weight' in style_dict and style_dict['font-weight'] in ['bold', '700']:
                run.bold = True
            if 'font-style' in style_dict and style_dict['font-style'] == 'italic':
                run.italic = True
            if 'color' in style_dict:
                run.font.color.rgb = parse_color(style_dict['color'])
            
            # Apply classes
            if 'text-underline' in classes:
                run.underline = True
            if 'text-strikethrough' in classes:
                run.font.strike = True
            if 'text-subscript' in classes:
                run.font.subscript = True
            if 'text-superscript' in classes:
                run.font.superscript = True
            if 'text-highlight' in classes:
                run.font.highlight_color = 7  # Yellow highlight
        elif child.name in ['b', 'strong']:
            run = para.add_run(child.get_text())
            run.bold = True
        elif child.name in ['i', 'em']:
            run = para.add_run(child.get_text())
            run.italic = True
        elif child.name == 'u':
            run = para.add_run(child.get_text())
            run.underline = True

def process_html_table_to_docx(doc, table_element):
    """Convert HTML table to Word table"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count columns
    cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    
    # Process each row
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            if col_idx < cols:
                # Clear default paragraph
                word_cell = table.cell(row_idx, col_idx)
                word_cell.paragraphs[0].clear()
                
                # Process cell content
                for element in cell.children:
                    if element.name == 'p':
                        process_html_paragraph_to_docx(word_cell, element)
                    elif isinstance(element, str) and element.strip():
                        para = word_cell.add_paragraph()
                        para.add_run(element.strip())

def process_html_list_to_docx(doc, list_element):
    """Convert HTML list to Word list"""
    items = list_element.find_all('li')
    for item in items:
        para = doc.add_paragraph(style='List Bullet' if list_element.name == 'ul' else 'List Number')
        process_inline_content(para, item)

def parse_style_string(style_str):
    """Parse CSS style string into dictionary"""
    style_dict = {}
    for item in style_str.split(';'):
        if ':' in item:
            key, value = item.split(':', 1)
            style_dict[key.strip()] = value.strip()
    return style_dict

def parse_pt_value(value):
    """Parse point value from CSS string"""
    match = re.search(r'(\d+(?:\.\d+)?)', value)
    if match:
        return float(match.group(1))
    return 0

def parse_color(color_str):
    """Parse color string to RGBColor"""
    if color_str.startswith('#'):
        # Hex color
        hex_str = color_str.lstrip('#')
        if len(hex_str) == 3:
            hex_str = ''.join([c*2 for c in hex_str])
        return RGBColor(
            int(hex_str[0:2], 16),
            int(hex_str[2:4], 16),
            int(hex_str[4:6], 16)
        )
    elif color_str.startswith('rgb'):
        # RGB color
        matches = re.findall(r'\d+', color_str)
        if len(matches) >= 3:
            return RGBColor(int(matches[0]), int(matches[1]), int(matches[2]))
    # Default black
    return RGBColor(0, 0, 0)

@app.route('/download-document', methods=['POST'])
def download_document():
    temp_file = None
    try:
        data = request.get_json()
        filename = data.get('filename')
        content = data.get('content')
        
        if not filename:
            return jsonify({'error': 'Filename is required'}), 400
            
        # Create a new document with the content
        doc = Document()
        
        # Parse HTML content
        soup = BeautifulSoup(content, 'html.parser')
        
        # Find all elements (paragraphs, headings, tables, lists)
        elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'table', 'ul', 'ol'])
        
        for element in elements:
            if element.name == 'table':
                # Process table
                process_html_table_to_docx(doc, element)
            elif element.name in ['ul', 'ol']:
                # Process list
                process_html_list_to_docx(doc, element)
            else:
                # Process paragraph or heading
                process_html_paragraph_to_docx(doc, element)
        
        # Generate a unique temporary file path
        temp_dir = tempfile.gettempdir()
        temp_file = os.path.join(temp_dir, f'temp_{int(time.time())}_{filename}')
        
        # Save the document to the temporary file
        doc.save(temp_file)
        
        # Create response with cleanup callback
        def generate():
            try:
                with open(temp_file, 'rb') as f:
                    yield from f
            finally:
                # Clean up after the file has been sent
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except Exception as e:
                    logger.error(f"Error cleaning up temporary file: {str(e)}")
        
        # Return file as a stream with cleanup
        return Response(
            generate(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"'
            }
        )
                
    except Exception as e:
        logger.error(f"Error downloading document: {str(e)}")
        # Clean up if error occurs before sending
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except Exception as cleanup_error:
                logger.error(f"Error cleaning up temporary file after error: {str(cleanup_error)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.chat_histories = {}
    app.run(host='0.0.0.0', port=8090, debug=True)