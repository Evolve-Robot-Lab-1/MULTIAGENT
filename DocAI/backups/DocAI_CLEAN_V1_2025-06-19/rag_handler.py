import os
import logging
import pickle
import shutil
from typing import Dict, Any, Optional
from langchain_community.document_loaders import TextLoader, PyPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain_huggingface import HuggingFaceEmbeddings

logger = logging.getLogger(__name__)

class RAGHandler:
    def __init__(self, uploads_dir: str = "uploads"):
        """Initialize with HuggingFace embeddings and uploads directory"""
        logger.info("Initializing RAGHandler with HuggingFace embeddings")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
        self.uploads_dir = uploads_dir
        self.embeddings_dir = os.path.join(uploads_dir, "embeddings")
        self.documents_dir = os.path.join(uploads_dir, "documents")
        self.vector_store = None
        self.document_store = {}
        self._initialize_directories()
        self._initialize_vector_store()

    def _initialize_directories(self):
        """Create necessary directories if they don't exist"""
        os.makedirs(self.embeddings_dir, exist_ok=True)
        os.makedirs(self.documents_dir, exist_ok=True)
        logger.info(f"Initialized directories: {self.embeddings_dir}, {self.documents_dir}")

    def _get_embedding_path(self, filename: str) -> str:
        """Get the path for storing embeddings for a specific file"""
        base_name = os.path.splitext(filename)[0]
        return os.path.join(self.embeddings_dir, f"{base_name}_embeddings.pkl")

    def _get_document_path(self, filename: str) -> str:
        """Get the path for storing the document"""
        return os.path.join(self.documents_dir, filename)

    def _initialize_vector_store(self):
        """Initialize vector store from existing embeddings"""
        try:
            # Check for any existing embeddings
            embedding_files = [f for f in os.listdir(self.embeddings_dir) if f.endswith('_embeddings.pkl')]
            if embedding_files:
                # Load the first embedding file to initialize the vector store
                first_embedding = os.path.join(self.embeddings_dir, embedding_files[0])
                with open(first_embedding, 'rb') as f:
                    self.vector_store = pickle.load(f)
                
                # Add other embeddings to the vector store
                for emb_file in embedding_files[1:]:
                    with open(os.path.join(self.embeddings_dir, emb_file), 'rb') as f:
                        vs = pickle.load(f)
                        if vs and hasattr(vs, 'docstore'):
                            self.vector_store.merge_from(vs)
                
                # Update document store
                for filename in os.listdir(self.documents_dir):
                    base_name = os.path.splitext(filename)[0]
                    if any(f.startswith(base_name) for f in embedding_files):
                        self.document_store[filename] = {
                            'path': self._get_document_path(filename)
                        }
                
                logger.info("Loaded existing vector store from embeddings directory")
            else:
                logger.info("No existing embeddings found")
                self.vector_store = None
        except Exception as e:
            logger.error(f"Error loading vector store: {str(e)}")
            self.vector_store = None

    def _get_document_loader(self, file_path: str):
        """Enhanced document loader with better formatting preservation"""
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            return PyPDFLoader(file_path)
        elif ext in ['doc', 'docx']:
            # Custom Word document loader with formatting preservation
            return self._enhanced_docx_loader(file_path)
        else:
            return TextLoader(file_path)

    def _enhanced_docx_loader(self, file_path: str):
        """Custom loader for Word documents that preserves exact formatting"""
        
        class EnhancedDocxLoader:
            def __init__(self, file_path):
                self.file_path = file_path
            
            def load(self):
                try:
                    from docx import Document as DocxDocument
                    from langchain_core.documents import Document as LangchainDocument
                    
                    # Load document with python-docx for better formatting access
                    doc = DocxDocument(self.file_path)
                    
                    # Extract content with formatting preserved
                    full_content = []
                    
                    # Process paragraphs and tables in order
                    for element in doc.element.body:
                        if element.tag.endswith('p'):
                            # Find matching paragraph
                            for para in doc.paragraphs:
                                if para._element == element:
                                    para_text = self._extract_paragraph_with_formatting(para)
                                    if para_text.strip():
                                        full_content.append(para_text)
                                    break
                        
                        elif element.tag.endswith('tbl'):
                            # Find matching table
                            for table in doc.tables:
                                if table._element == element:
                                    table_text = self._extract_table_text(table)
                                    if table_text.strip():
                                        full_content.append(table_text)
                                    break
                    
                    # Create document with enhanced metadata
                    content = '\n\n'.join(full_content)
                    
                    return [LangchainDocument(
                        page_content=content,
                        metadata={
                            'source': self.file_path,
                            'type': 'docx',
                            'document_type': 'word_document',
                            'has_formatting': True
                        }
                    )]
                    
                except Exception as e:
                    logger.error(f"Error in enhanced docx loader: {str(e)}")
                    # Fallback to standard loader
                    try:
                        fallback_loader = Docx2txtLoader(self.file_path)
                        documents = fallback_loader.load()
                        for doc in documents:
                            doc.metadata.update({
                                'source': self.file_path,
                                'type': 'docx',
                                'document_type': 'word_document'
                            })
                        return documents
                    except:
                        from langchain_core.documents import Document as LangchainDocument
                        return [LangchainDocument(
                            page_content="", 
                            metadata={'source': self.file_path, 'type': 'docx', 'error': str(e)}
                        )]
            
            def _extract_paragraph_with_formatting(self, para):
                """Extract paragraph text with formatting indicators"""
                if not para.text.strip():
                    return ""
                
                # Add style indicators
                prefix = ""
                if para.style and para.style.name:
                    if 'Heading' in para.style.name:
                        level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                        prefix = f"[HEADING{level}] "
                
                # Add alignment indicator if not default
                if para.alignment and para.alignment != 0:  # 0 is left align
                    alignments = {1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY'}
                    if para.alignment in alignments:
                        prefix += f"[{alignments[para.alignment]}] "
                
                return prefix + para.text
            
            def _extract_table_text(self, table):
                """Extract table content as structured text"""
                rows_text = []
                rows_text.append("[TABLE]")
                
                for row in table.rows:
                    cells_text = []
                    for cell in row.cells:
                        cell_content = ' '.join(p.text for p in cell.paragraphs if p.text.strip())
                        cells_text.append(cell_content)
                    
                    if any(cells_text):
                        rows_text.append(" | ".join(cells_text))
                
                rows_text.append("[/TABLE]")
                return '\n'.join(rows_text)
        
        return EnhancedDocxLoader(file_path)

    def process_document(self, file_path: str, filename: str) -> bool:
        try:
            # Save original document
            document_path = self._get_document_path(filename)
            shutil.copy2(file_path, document_path)
            
            # Load and extract text with formatting
            loader = self._get_document_loader(document_path)
            documents = loader.load()
            
            # Enhanced text splitting that preserves formatting markers
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,  # Increased chunk size to better preserve context
                chunk_overlap=400,
                separators=["\n\n", "\n", " ", ""],
                keep_separator=True,
                add_start_index=True,
                strip_whitespace=False  # Preserve whitespace for formatting
            )
            
            # Custom preprocessing to handle formatting markers
            def preprocess_text(text):
                # Preserve formatting markers during splitting
                markers = ['[SECTION_PROPS]', '[STYLES]', '[PARAGRAPH]', '[TABLE]']
                for marker in markers:
                    text = text.replace(marker, f"\n{marker}\n")
                return text

            # Apply preprocessing and split documents
            processed_docs = []
            for doc in documents:
                doc.page_content = preprocess_text(doc.page_content)
                processed_docs.extend(text_splitter.split_documents([doc]))

            # Create embeddings 
            new_vector_store = FAISS.from_documents(
                processed_docs,
                self.embeddings
            )
            
            if not self.vector_store:
                self.vector_store = new_vector_store
            else:
                self.vector_store.merge_from(new_vector_store)

            # Save embeddings
            embedding_path = self._get_embedding_path(filename)
            with open(embedding_path, 'wb') as f:
                pickle.dump(new_vector_store, f)

            self.document_store[filename] = {
                'path': document_path,
                'has_formatting': True
            }

            return True

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return False

    def remove_document(self, filename: str) -> bool:
        try:
            if filename in self.document_store:
                # Remove document file
                document_path = self._get_document_path(filename)
                if os.path.exists(document_path):
                    os.remove(document_path)

                # Remove embeddings file
                embedding_path = self._get_embedding_path(filename)
                if os.path.exists(embedding_path):
                    os.remove(embedding_path)

                # Remove from document store
                del self.document_store[filename]

                # Reinitialize vector store from remaining embeddings
                self.vector_store = None
                self._initialize_vector_store()

                logger.info(f"Successfully removed document and embeddings for {filename}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error removing document {filename}: {str(e)}")
            return False

    def query_documents(self, query: str, llm) -> Optional[str]:
        try:
            if not self.vector_store:
                return None

            # Generic prompt for any document type
            prompt = ChatPromptTemplate.from_template("""
                Answer the question based on the provided context from the uploaded documents.
                
                <context>
                {context}
                </context>
                
                Question: {input}
                
                Provide a clear and accurate response based on the information found in the documents.
                If the information is not available in the context, say "I don't have that information in the uploaded documents."
            """)

            document_chain = create_stuff_documents_chain(llm, prompt)
            
            # Use retrieval without specific filtering
            retriever = self.vector_store.as_retriever(
                search_kwargs={
                    "k": 5,
                    "fetch_k": 10
                }
            )
            
            retrieval_chain = create_retrieval_chain(retriever, document_chain)
            response = retrieval_chain.invoke({"input": query})
            
            return response.get("answer", None)

        except Exception as e:
            logger.error(f"Error querying documents: {str(e)}")
            return None

    def get_status(self) -> Dict[str, Any]:
        return {
            "has_documents": self.has_documents(),
            "document_count": len(self.document_store),
            "documents": list(self.document_store.keys())
        }

    def has_documents(self) -> bool:
        return self.vector_store is not None and len(self.document_store) > 0